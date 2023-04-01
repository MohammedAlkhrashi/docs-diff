from datetime import date
import argparse

import os
import os.path as osp
from docx import Document
from reportlab.lib.pagesizes import letter

import docx2txt

from reportlab.pdfgen import canvas
from reportlab.lib.units import inch
from docx.enum.text import WD_COLOR_INDEX
import myers
from reportlab.lib.colors import Color, black, red, green

# def compare_docx(path_old, path_new):
#     DOC1 = aw.Document(path_old)
#     DOC2 = aw.Document(path_new)

#     options = aw.comparing.CompareOptions()
#     options.ignore_formatting = True
#     options.ignore_headers_and_footers = True
#     options.ignore_case_changes = True
#     options.ignore_tables = True
#     options.ignore_fields = True
#     options.ignore_comments = True
#     options.ignore_textboxes = True
#     options.ignore_footnotes = True

#     # DOC1 will contain changes as revisions after comparison
#     DOC1.compare(DOC2, "user", date.today(), options)

#     if DOC1.revisions.count > 0:
#         # Save resultant file as PDF
#         DOC1.save("compared.pdf", aw.SaveFormat.PDF)
#     else:
#         print("Documents are equal")


# def compare_pdf(path_old, path_new):
#     # Load PDF files
#     PDF1 = aw.Document(path_old)
#     PDF2 = aw.Document(path_new)

#     def convert_path_pdf_to_docx(path):
#         return osp.splitext(path)[0] + ".docx"

#     path_old_docx = convert_path_pdf_to_docx(path_old)
#     path_new_docx = convert_path_pdf_to_docx(path_new)

#     # Convert PDF files to Word format
#     PDF1.save(path_old_docx, aw.SaveFormat.DOCX)
#     PDF1.save(path_new_docx, aw.SaveFormat.DOCX)
#     compare_docx(path_old_docx, path_new_docx)


# def compare_docx(old_path, new_path):
#     old_doc = Document(old_path)
#     new_doc = Document(new_path)

#     old_text = "\n".join([para.text for para in old_doc.paragraphs])
#     new_text = "\n".join([para.text for para in new_doc.paragraphs])

#     # d = difflib.Differ()
#     # diff = list(d.compare(old_text.splitlines(), new_path.splitlines()))
#     # diff = list(difflib.unified_diff(old_text.splitlines(), new_text.splitlines()))
#     diff = difflib.ndiff(old_text.splitlines(), new_text.splitlines())

#     pdf = canvas.Canvas("diff.pdf")
#     pdf.setFont("Helvetica", 10)

#     pdf.drawString(0.5 * inch, 10 * inch, old_text)
#     pdf.drawString(0.5 * inch, 9 * inch, new_text)

#     # Loop through the list of differences and add annotations to the PDF file
#     for line in diff:
#         print(line)
#         if line.startswith("+"):
#             # Added text
#             pdf.setFillColorRGB(0, 1, 0)  # green
#             pdf.drawString(0.5 * inch, 8 * inch, line[2:])
#         elif line.startswith("-"):
#             # Deleted text
#             pdf.setStrokeColorRGB(1, 0, 0)  # red
#             pdf.line(0.5 * inch, 7.8 * inch, 8 * inch, 7.8 * inch)
#             pdf.drawString(0.5 * inch, 7.5 * inch, line[2:])
#         elif line.startswith("?"):
#             # Replaced text
#             pdf.setFillColorRGB(1, 1, 0)  # yellow
#             pdf.drawString(0.5 * inch, 7 * inch, line[2:])

#     # Save the new PDF file with the annotated changes
#     pdf.save()


def compare_docx(old_path, new_path):
    doc1_text = docx2txt.process(old_path)
    doc2_text = docx2txt.process(new_path)

    diff = myers.diff(doc1_text.splitlines(), doc2_text.splitlines())

    c = canvas.Canvas("diff.pdf", pagesize=(8.5 * inch, 11 * inch))
    c.setFont("Courier", 10)
    x, y = 1 * inch, 10.5 * inch

    red_color = Color(1, 0, 0, alpha=0.8)
    green_color = Color(0, 1, 0, alpha=0.8)

    for action, line in diff:
        print({action: line})
        if action == "k":
            c.setFillColor(black)
            c.drawString(x, y, line)
            y -= 12
        elif action == "i":
            c.setFillColor(green_color)
            c.drawString(x, y, "+" + line)
            y -= 12
        elif action == "r":
            c.setFillColor(red_color)
            c.drawString(x, y, "-" + line)
            y -= 12
        elif action == "o":
            pass

    c.save()


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--old_path", type=str, help="path to old file", default="./old_file.docx"
    )
    parser.add_argument(
        "--new_path", type=str, help="path to new file", default="./new_file.docx"
    )
    args = parser.parse_args()

    old_path = args.old_path
    new_path = args.new_path
    compare_docx(old_path, new_path)


if __name__ == "__main__":
    main()
